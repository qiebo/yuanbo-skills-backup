import argparse
import json
import re
import shutil
from datetime import datetime, timedelta, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


TZ = timezone(timedelta(hours=8))


def default_template_path():
    root = Path(__file__).resolve().parents[3]
    preferred = root / "输入" / "XXXX学校的生涯翼站使用分析报告.docx"
    if preferred.exists():
        return preferred
    return Path(__file__).resolve().parents[1] / "assets" / "templates" / "yizhan-report-template.docx"


def load_json(path):
    return json.loads(Path(path).read_text(encoding="utf-8"))


def num(value):
    if value is None or value == "":
        return 0
    if isinstance(value, (int, float)):
        return value
    text = str(value).replace(",", "").replace("%", "")
    try:
        n = float(text)
    except ValueError:
        return 0
    return n / 100 if str(value).strip().endswith("%") else n


def fmt_num(value, decimals=0):
    if value is None or value == "":
        return "未获取"
    if isinstance(value, str):
        return value
    if isinstance(value, float) and not value.is_integer():
        return f"{value:,.{decimals}f}"
    return f"{int(value):,}"


def fmt_pct(value, decimals=1):
    if value is None or value == "":
        return "未获取"
    if isinstance(value, str):
        return value if value.endswith("%") else value
    return f"{value * 100:.{decimals}f}%"


def fmt_period(period):
    start = period.get("start_month", "")
    end = period.get("end_month", "")
    return f"{fmt_month(start)} — {fmt_month(end)}"


def fmt_month(month):
    if not month or "-" not in str(month):
        return str(month or "")
    year, raw_month = str(month).split("-", 1)
    return f"{int(year)} 年 {int(raw_month)} 月"


def core(registry, name, default=None):
    for row in registry.get("core_metrics", []):
        if row.get("指标名称") == name:
            return row.get("数值", default)
    return default


def metric(registry, module, name, default=None):
    for row in registry.get("module_metrics", []):
        if row.get("模块") == module and row.get("指标") == name:
            return row.get("数值", default)
    return default


def content(registry, module=None, dimension=None):
    rows = registry.get("content_and_grade_stats", [])
    if module:
        rows = [r for r in rows if r.get("模块") == module]
    if dimension:
        rows = [r for r in rows if r.get("维度") == dimension]
    return rows


def grade_order(value):
    text = str(value)
    for idx, grade in enumerate(["高一", "高二", "高三"], start=1):
        if grade in text:
            return idx
    return 99


def display_grade(value):
    text = str(value or "")
    if text.startswith("2025-"):
        return text.replace("2025-", "") + " (2025级)"
    return text


def grade_values(rows):
    values = []
    for row in rows:
        obj = row.get("对象")
        if obj and obj not in values:
            values.append(obj)
    return sorted(values, key=grade_order)


def lookup_content(registry, module, dimension, grade, category, default="0"):
    for row in content(registry, module, dimension):
        if row.get("对象") == grade and row.get("分类或指标") == category:
            return row.get("数值", default)
    return default


def top_questions(registry, limit=8):
    blocked = ["AV", "不当", "黄色", "色情"]
    rows = sorted(content(registry, "AI生涯规划", "高频提问"), key=lambda r: num(r.get("数值")), reverse=True)
    result = []
    for row in rows:
        text = str(row.get("分类或指标") or "").strip()
        if text and not any(term.lower() in text.lower() for term in blocked):
            result.append(text)
        if len(result) >= limit:
            break
    while len(result) < limit:
        result.append("围绕专业选择、升学路径、职业方向和学习规划的常见问题")
    return result


def first_run_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = str(text)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(str(text))


def set_paragraph_pagination(paragraph, keep_next=None, keep_together=None, page_break_before=None):
    fmt = paragraph.paragraph_format
    if keep_next is not None:
        fmt.keep_with_next = keep_next
    if keep_together is not None:
        fmt.keep_together = keep_together
    if page_break_before is not None:
        fmt.page_break_before = page_break_before


def remove_repeated_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    for tbl_header in list(tr_pr.findall(qn("w:tblHeader"))):
        tr_pr.remove(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def apply_table_pagination_rules(doc):
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            remove_repeated_table_header(row)
            keep_row_together(row)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_paragraph_pagination(
                        paragraph,
                        keep_next=row_idx < len(table.rows) - 1,
                        keep_together=True,
                    )

    for idx in [69, 70, 77, 78]:
        if idx <= len(doc.paragraphs):
            set_paragraph_pagination(doc.paragraphs[idx - 1], keep_next=True, keep_together=True)

    if len(doc.paragraphs) >= 77:
        set_paragraph_pagination(doc.paragraphs[76], page_break_before=True)


def remove_template_time_labels(doc):
    patterns = [
        r"（20\d{2}年\d{1,2}月数据）",
        r"（20\d{2}年数据）",
        r"（20\d{2}年\d{1,2}月）",
        r"\(20\d{2}年\d{1,2}月数据\)",
        r"\(20\d{2}年数据\)",
        r"\(20\d{2}年\d{1,2}月\)",
        r"^20\d{2}年\d{1,2}月",
    ]

    def clean_paragraph(paragraph):
        text = paragraph.text
        cleaned = text
        for pattern in patterns:
            cleaned = re.sub(pattern, "", cleaned)
        if cleaned != text:
            first_run_text(paragraph, cleaned)

    for paragraph in doc.paragraphs:
        clean_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    clean_paragraph(paragraph)


def cell_text(cell, text):
    if not cell.paragraphs:
        cell.add_paragraph()
    first_run_text(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        first_run_text(paragraph, "")


def paragraph(doc, index_1based, text):
    first_run_text(doc.paragraphs[index_1based - 1], text)


def table_cell(doc, table_no, row_no, col_no, text):
    cell_text(doc.tables[table_no - 1].rows[row_no - 1].cells[col_no - 1], text)


def set_row(doc, table_no, row_no, values):
    for col, value in enumerate(values, start=1):
        table_cell(doc, table_no, row_no, col, value)


def stage_groups(registry):
    months = {row.get("month"): row for row in registry.get("excel_monthly_usage", [])}
    backend = {row.get("month"): row for row in registry.get("website_backend_record_trend", [])}
    groups = [
        ("2025-12 — 2026-01", ["2025-12", "2026-01"], "集中启动期"),
        ("2026-02 — 2026-03", ["2026-02", "2026-03"], "开学深化期"),
        ("2026-04 — 2026-05", ["2026-04", "2026-05"], "常态应用期"),
    ]
    result = []
    for label, keys, stage in groups:
        excel_rows = [months[k] for k in keys if k in months]
        backend_rows = [backend[k] for k in keys if k in backend]
        active_peak = max([num(r.get("total_user_count")) for r in excel_rows] or [0])
        backend_total = sum(num(r.get("total_backend_records")) for r in backend_rows)
        login_total = sum(num(r.get("total_login_count")) for r in excel_rows)
        career = sum(num(r.get("ai_career_data_count")) for r in excel_rows)
        partner = sum(num(r.get("ai_partner_data_count")) for r in excel_rows)
        bottle = sum(num(r.get("message_bottle_total")) for r in excel_rows)
        interview = sum(num(r.get("ai_interview_total")) for r in excel_rows)
        result.append(
            {
                "label": label,
                "keys": keys,
                "stage": stage,
                "active_peak": active_peak,
                "backend_total": backend_total,
                "login_total": login_total,
                "career": career,
                "partner": partner,
                "bottle": bottle,
                "interview": interview,
            }
        )
    return result


def max_month(registry, field):
    rows = registry.get("excel_monthly_usage", [])
    if not rows:
        return None
    return max(rows, key=lambda r: num(r.get(field)))


def trend_analysis(registry):
    stages = stage_groups(registry)
    peak_active = max_month(registry, "total_user_count")
    peak_bottle = max_month(registry, "message_bottle_total")
    return [
        (
            f"{stages[0]['label']}：{stages[0]['stage']}",
            "首个完整活跃窗口出现在 1 月，登录人次和模块使用同步抬升，说明平台并不是被动浏览型工具，而是在学生遇到方向选择、情绪表达和升学准备问题时被主动调用。对学校而言，这一阶段的关键不是追求全校铺开，而是识别最先使用平台的学生群体和真实问题类型，为后续分年级运营确定入口。",
        ),
        (
            f"{stages[1]['label']}：{stages[1]['stage']}",
            "3 月形成全周期活跃峰值，生涯规划、心语伙伴和漂流瓶同时放量，表明开学后的学习节奏变化会把生涯困惑、心理表达和同伴互动需求集中释放。这个阶段更适合承接主题班会、生涯课和心理支持活动，而不是只做后台数据展示。",
        ),
        (
            f"{stages[2]['label']}：{stages[2]['stage']}",
            "4-5 月总体活跃回落，但漂流瓶仍保持较高互动，说明学生从“问规划、问方向”的任务型使用，逐步转向“表达情绪、回应同伴”的日常型使用。后续运营应把漂流瓶作为校园氛围和学生状态的观察窗口，同时保持内容巡检和正向引导。",
        ),
        (
            "整体趋势判断",
            f"本周期活跃峰值出现在 {peak_active.get('month') if peak_active else '未获取'}，漂流瓶高点出现在 {peak_bottle.get('month') if peak_bottle else '未获取'}。二者不完全重合，说明平台价值不是单一流量增长，而是不同功能在不同学生场景中承担不同任务：生涯规划承接方向决策，心语伙伴承接心理表达，漂流瓶承接同伴互助，模拟面试承接升学训练。",
        ),
    ]


def fill_stage_table(doc, registry, student_accounts):
    stages = stage_groups(registry)
    headers = ["维度"] + [s["label"] for s in stages]
    set_row(doc, 3, 1, headers)
    set_row(doc, 3, 2, ["阶段定位"] + [s["stage"] for s in stages])
    set_row(doc, 3, 3, [
        "核心特征",
        "1月形成首个集中活跃窗口",
        "3月多模块同步放量",
        "表达互动转向日常化",
    ])
    set_row(doc, 3, 4, [
        "用户行为",
        "从试用进入主动提问",
        "从咨询扩展到表达与互助",
        "从集中使用转向持续参与",
    ])
    set_row(doc, 3, 5, [
        "学期活跃覆盖率",
        *[f"峰值 {fmt_num(s['active_peak'])} 人（{fmt_pct(s['active_peak'] / student_accounts if student_accounts else None, 1)}）" for s in stages],
    ])


def grade_scene(grade):
    if "高一" in grade:
        return "学业规划起步、生活适应"
    if "高二" in grade:
        return "选科决策、生涯深度探索（需求峰值期）"
    if "高三" in grade:
        return "志愿填报、升学情绪疏导"
    return "生涯探索与学习规划"


def career_category_implication(category):
    mapping = {
        "院校专业类": "适合转化为专业认知、院校层次和选科关联的专题答疑，减少学生只在临近升学节点才集中求助",
        "自我认知与探索类": "说明学生需要先厘清兴趣、能力和价值取向，学校可把测评解读与个别咨询前置到生涯启蒙阶段",
        "学业规划类": "反映学生在学习节奏、目标拆解和执行方法上需要支持，适合与班主任学业指导和月度复盘结合",
        "职业行业类": "说明学生开始把专业选择和真实职业连接，适合引入校友、行业案例或职业访谈材料",
        "咨询工具类": "表明学生正在学习如何使用平台解决具体问题，后续应提供更清晰的使用场景和提问示例",
        "生活闲聊": "提示该年级把平台当作低压力表达入口，运营上应从轻量陪伴逐步引导到学业规划和自我探索任务",
    }
    return mapping.get(category, "建议把该类问题整理成年级常见问题清单，供生涯课、班会和个别咨询复用")


def partner_factor_implication(grade, factor):
    grade_text = str(grade or "")
    if factor == "人际关系":
        if "高一" in grade_text:
            return "更应放在新环境适应、同伴融入和班级规则建立上处理，帮助学生把入学初期的社交不确定感转化为可求助的具体问题"
        if "高二" in grade_text:
            return "可结合选科后学习圈层变化、合作学习和同伴支持小组，提前处理关系摩擦对学习状态的牵连"
        if "高三" in grade_text:
            return "更可能与升学压力下的同伴比较、沟通敏感和情绪承压叠加，适合在毕业年级开展关系调适和压力表达支持"
    if factor == "自信心":
        if "高一" in grade_text:
            return "应结合入学适应、优势发现和小目标达成，让学生尽早形成可被看见的积极反馈"
        if "高二" in grade_text:
            return "应结合阶段考试反馈、优势科目识别和目标拆解，帮助学生把短期挫败感转化为可执行的改进计划"
        if "高三" in grade_text:
            return "需要和志愿定位、模拟结果复盘及压力管理结合，避免一次性成绩波动被学生解释为整体能力否定"

    mapping = {
        "学业压力": "应优先承接学习节奏、目标分解和压力管理支持，避免压力长期沉淀为情绪问题",
        "考试焦虑": "适合在考试前后设置情绪调适、复盘方法和求助渠道，让学生知道焦虑可以被具体处理",
        "个人形象": "更适合通过青春期自我认同、同伴比较和身体意象教育来承接，避免简单归因为个体敏感",
        "家庭关系": "需要班主任和心理老师谨慎复核，重点关注沟通压力、支持系统和学生可获得的现实帮助",
    }
    return mapping.get(factor, "建议作为线索进入人工复核，而不是直接把后台比例当作结论")


def partner_state_implication(grade, concern, special):
    grade_text = str(grade or "")
    special_rate = num(special)
    if special_rate >= 0.1:
        if "高二" in grade_text:
            return "特别关注占比较高，应优先建立小样本复核台账，并由班主任与心理老师共同判断是否存在持续性压力源"
        return "特别关注比例已经不宜只做趋势观察，应建立名单复核、班主任沟通和心理老师跟进的闭环"
    if concern >= 0.6:
        if "高一" in grade_text:
            return "更像入学适应和同伴融入阶段的集中表达，建议把班级观察、适应性班会和个别谈话结合起来复核"
        if "高三" in grade_text:
            return "应优先结合升学压力、考试节奏和近期情绪波动做分层判断，避免把高比例简单解释成普遍风险"
        return "需关注样本占比较高，应把后续工作放在分层复核、重点样本跟进和资源匹配上"
    if concern >= 0.3:
        return "存在一定关注样本，适合纳入常态观察池，并结合线下表现判断是否需要进一步支持"
    return "当前关注比例相对有限，但仍应保留持续观察，避免低频表达被忽略"


def fill_tables(doc, registry):
    school = registry.get("school_name", "学校")
    student_accounts = core(registry, "学生账号总数")
    served = core(registry, "累计服务学生")
    total_interactions = core(registry, "累计交互总量")
    coverage = core(registry, "去重服务覆盖率")
    if coverage is None:
        coverage = core(registry, "全校覆盖率")
    active_peak = core(registry, "本学期活跃峰值")
    highest_module = core(registry, "全平台覆盖率最高模块")

    career_users = metric(registry, "AI生涯规划", "覆盖学生")
    career_uses = metric(registry, "AI生涯规划", "累计使用次数")
    career_avg = career_uses / career_users if career_users else None
    partner_dialogs = metric(registry, "AI心语伙伴", "累计对话")
    partner_users = metric(registry, "AI心语伙伴", "覆盖学生")
    partner_avg = partner_dialogs / partner_users if partner_users else None
    bottle_posts = metric(registry, "漂流瓶", "发布数量")
    bottle_replies = metric(registry, "漂流瓶", "回复数量")
    bottle_total = metric(registry, "漂流瓶", "累计收发") or ((bottle_posts or 0) + (bottle_replies or 0))
    bottle_reply_rate = metric(registry, "漂流瓶", "回复率")
    bottle_unique = metric(registry, "漂流瓶", "发瓶学生去重数")
    interview_users = metric(registry, "AI模拟面试", "服务学生")
    interview_records = metric(registry, "AI模拟面试", "累计使用人次")
    interview_reports = metric(registry, "AI模拟面试", "生成报告")
    interview_depth = metric(registry, "AI模拟面试", "深度使用率")

    table_cell(doc, 1, 1, 2, fmt_period(registry.get("data_period", {})))
    table_cell(doc, 1, 2, 2, "AI 生涯规划 / AI 心语伙伴 / 漂流瓶 / AI 模拟面试")
    table_cell(doc, 1, 3, 2, f"{fmt_num(student_accounts)} 人")
    now = datetime.now(TZ)
    table_cell(doc, 1, 4, 2, f"{now.year} 年 {now.month} 月 {now.day} 日")

    set_row(doc, 2, 2, [f"{fmt_num(served)} 人", f"{fmt_num(total_interactions)} 条", fmt_pct(coverage, 1), f"{fmt_num(active_peak)} 人"])
    fill_stage_table(doc, registry, student_accounts)
    set_row(doc, 4, 2, [f"{fmt_num(career_uses)} 条", f"{fmt_num(career_users)} 人", f"{career_avg:.1f} 次" if career_avg else "未获取"])

    for idx, question in enumerate(top_questions(registry, 8), start=2):
        set_row(doc, 5, idx, [str(idx - 1), question])

    grades = grade_values(content(registry, "AI生涯规划", "年级使用占比"))
    for row_no, grade in enumerate(grades[:3], start=2):
        rate = lookup_content(registry, "AI生涯规划", "年级使用占比", grade, "使用人数占比", "0%")
        set_row(doc, 6, row_no, [display_grade(grade), rate, grade_scene(grade)])

    content_grades = grade_values(content(registry, "AI生涯规划", "年级互动内容分类"))
    for row_no, grade in enumerate(content_grades[:3], start=2):
        set_row(doc, 7, row_no, [
            display_grade(grade),
            fmt_num(lookup_content(registry, "AI生涯规划", "年级互动内容分类", grade, "院校专业类")),
            fmt_num(lookup_content(registry, "AI生涯规划", "年级互动内容分类", grade, "自我认知与探索类")),
            fmt_num(lookup_content(registry, "AI生涯规划", "年级互动内容分类", grade, "学业规划类")),
            fmt_num(lookup_content(registry, "AI生涯规划", "年级互动内容分类", grade, "职业行业类")),
            fmt_num(lookup_content(registry, "AI生涯规划", "年级互动内容分类", grade, "咨询工具类")),
            fmt_num(lookup_content(registry, "AI生涯规划", "年级互动内容分类", grade, "生活闲聊")),
            fmt_num(lookup_content(registry, "AI生涯规划", "年级互动内容分类", grade, "其他")),
        ])

    set_row(doc, 8, 2, [f"{fmt_num(partner_dialogs)} 条", f"{fmt_num(partner_users)} 人", f"{partner_avg:.1f} 条" if partner_avg else "未获取"])

    factor_grades = grade_values(content(registry, "AI心语伙伴", "心理困扰因子年级分布"))
    factor_names = ["学业压力", "考试焦虑", "个人形象", "家庭关系", "人际关系", "自信心"]
    for row_no, grade in enumerate(factor_grades[:3], start=2):
        set_row(doc, 9, row_no, [display_grade(grade)] + [lookup_content(registry, "AI心语伙伴", "心理困扰因子年级分布", grade, f, "0%") for f in factor_names])

    state_grades = grade_values(content(registry, "AI心语伙伴", "心理状态年级分布"))
    states = ["正常交流", "初步关注", "关怀提醒", "特别关注"]
    for row_no, grade in enumerate(state_grades[:3], start=2):
        values = [lookup_content(registry, "AI心语伙伴", "心理状态年级分布", grade, s, "0%") for s in states]
        concern = fmt_pct(sum(num(v) for v in values[1:]), 1)
        set_row(doc, 10, row_no, [display_grade(grade), *values, concern])

    set_row(doc, 11, 2, [f"{fmt_num(bottle_total)} 条", fmt_pct(bottle_reply_rate, 1), f"{fmt_num(bottle_unique)} 人", "最高" if highest_module == "漂流瓶" else str(highest_module or "")])
    set_row(doc, 12, 2, [f"{fmt_num(interview_users)} 人", f"{fmt_num(interview_records)} 人次", f"{fmt_num(interview_reports)} 份", fmt_pct(interview_depth, 1)])
    set_row(doc, 13, 2, ["漂流瓶回复率", f"{fmt_pct(bottle_reply_rate, 1)} — 平均每发出 1 条获得约 {(bottle_replies / bottle_posts):.1f} 条回复" if bottle_posts else "未获取"])
    set_row(doc, 13, 3, ["AI 心语伙伴人均对话", f"{partner_avg:.1f} 条 / 人 — 使用深度仍有提升空间" if partner_avg else "未获取"])
    set_row(doc, 13, 4, ["AI 生涯规划人均使用", f"{career_avg:.1f} 次 / 人 — 反复探索型功能" if career_avg else "未获取"])
    set_row(doc, 13, 5, ["用户行为成熟度", "已出现主动咨询、匿名表达、同伴回应和报告复盘四类行为"])
    set_row(doc, 14, 2, ["补位心理健康工作", "AI 心语伙伴可作为学生主动表达和风险线索初筛入口，但当前样本量仍需扩大；关键价值在于帮助学校更早发现需要人工复核的学生。"])
    set_row(doc, 14, 3, ["落地生涯教育实践", "高频问题集中在兴趣方向、专业选择、时间管理和志愿规则，适合转化为年级化生涯课程与专题答疑。"])
    set_row(doc, 14, 4, ["营造正向校园文化", "漂流瓶回复率高，说明同伴回应意愿强；学校可通过正向话题和内容巡检，把匿名表达转化为校园互助资源。"])
    set_row(doc, 14, 5, ["强化升学服务能力", "AI 模拟面试已形成一定报告转化，适合围绕系统已开放的模拟练习、智能评分、语音回听和改进建议组织训练与报告复盘。"])
    set_row(doc, 14, 6, ["提供数据化管理依据", "登记表把来源、口径和限制固定下来，便于后续按月复盘，不再依赖临时截图或人工描述。"])

    return {
        "school": school,
        "student_accounts": student_accounts,
        "served": served,
        "coverage": coverage,
        "career_users": career_users,
        "career_uses": career_uses,
        "career_avg": career_avg,
        "partner_dialogs": partner_dialogs,
        "partner_users": partner_users,
        "partner_avg": partner_avg,
        "bottle_total": bottle_total,
        "bottle_reply_rate": bottle_reply_rate,
        "bottle_posts": bottle_posts,
        "bottle_replies": bottle_replies,
        "bottle_unique": bottle_unique,
        "interview_users": interview_users,
        "interview_records": interview_records,
        "interview_reports": interview_reports,
        "interview_depth": interview_depth,
        "highest_module": highest_module,
    }


def insight_for_career(registry):
    grades = grade_values(content(registry, "AI生涯规划", "年级使用占比"))
    top_grade = None
    if grades:
        top_grade = max(grades, key=lambda g: num(lookup_content(registry, "AI生涯规划", "年级使用占比", g, "使用人数占比", "0%")))
    top_question = top_questions(registry, 1)[0]
    return (
        f"从高频问题看，学生最关心的不是简单了解职业信息，而是“{top_question}”这类自我定位与选择依据问题，说明平台正在承接学校线下生涯辅导中最耗时的一类个别化咨询。"
        + (f"{display_grade(top_grade)}使用占比最高，意味着当前生涯规划需求更集中在升学决策和专业选择阶段，后续可把高频问题沉淀为班会、讲座和个别咨询清单。" if top_grade else "")
    )


def insight_for_career_table(registry):
    grades = grade_values(content(registry, "AI生涯规划", "年级互动内容分类"))
    if not grades:
        return ["生涯规划内容分类数据不足，建议后续持续补充。"] * 3
    rows = []
    for grade in grades[:3]:
        cats = {
            "院校专业类": num(lookup_content(registry, "AI生涯规划", "年级互动内容分类", grade, "院校专业类")),
            "学业规划类": num(lookup_content(registry, "AI生涯规划", "年级互动内容分类", grade, "学业规划类")),
            "生活闲聊": num(lookup_content(registry, "AI生涯规划", "年级互动内容分类", grade, "生活闲聊")),
            "职业行业类": num(lookup_content(registry, "AI生涯规划", "年级互动内容分类", grade, "职业行业类")),
        }
        top = max(cats, key=cats.get)
        rows.append(f"{display_grade(grade)}的最高话题为{top}，{career_category_implication(top)}。")
    while len(rows) < 3:
        rows.append("各年级应按真实高频话题设计不同运营入口。")
    return rows


def insight_for_partner(registry):
    rows = content(registry, "AI心语伙伴", "心理困扰因子年级分布")
    if not rows:
        return ["心理困扰因子数据不足，建议后续补充。"] * 3
    grouped = {}
    for row in rows:
        grouped.setdefault(row.get("对象"), []).append(row)
    insights = []
    for grade, grade_rows in sorted(grouped.items(), key=lambda item: grade_order(item[0]))[:3]:
        top = max(grade_rows, key=lambda r: num(r.get("数值")))
        factor = top.get("分类或指标")
        insights.append(f"{display_grade(grade)}最突出的因子是{factor}（{top.get('数值')}），{partner_factor_implication(grade, factor)}。")
    while len(insights) < 3:
        insights.append("心理数据仅代表已互动样本，正式干预前需要教师结合现实表现复核。")
    return insights


def insight_for_partner_state(registry):
    rows = content(registry, "AI心语伙伴", "心理状态年级分布")
    grouped = {}
    for row in rows:
        grouped.setdefault(row.get("对象"), {})[row.get("分类或指标")] = row.get("数值")
    insights = []
    for grade, values in sorted(grouped.items(), key=lambda item: grade_order(item[0]))[:3]:
        concern = sum(num(values.get(k, 0)) for k in ["初步关注", "关怀提醒", "特别关注"])
        special = values.get("特别关注", "0%")
        insights.append(f"{display_grade(grade)}需关注比例为{fmt_pct(concern, 1)}，其中特别关注为{special}。{partner_state_implication(grade, concern, special)}。")
    while len(insights) < 3:
        insights.append("分级预警仅代表互动样本，不应外推为全校风险率。")
    return insights


def fill_paragraphs(doc, registry, values):
    school = values["school"]
    paragraph(doc, 2, school.replace("学校", "学  校") if len(school) <= 8 else school)
    paragraph(doc, 3, "生涯翼站使用情况学期分析报告")
    paragraph(doc, 8, f"本报告基于生涯翼站 SaaS 平台后台数据，以学期为单元对{school}学生使用行为、功能价值与教育成效进行深度分析，为学校生涯教育、心理健康建设与升学服务提供决策依据。")
    paragraph(
        doc,
        12,
        f"本周期累计服务学生 {fmt_num(values['served'])} 人，全校覆盖率 {fmt_pct(values['coverage'], 1)}，说明平台已形成有效使用样本但仍有扩面空间；{values['highest_module']}是使用贡献最高模块，漂流瓶回复率达到 {fmt_pct(values['bottle_reply_rate'], 1)}，反映学生表达与同伴回应需求明显；AI 模拟面试已生成 {fmt_num(values['interview_reports'])} 份报告，具备升学训练转化基础。",
    )
    paragraph(doc, 15, "平台已形成生涯规划、心理陪伴、校园互助、升学实战四位一体的服务框架，但当前重点应从“功能可用”转向“年级场景运营”")
    paragraph(doc, 16, f"AI 心语伙伴累计对话 {fmt_num(values['partner_dialogs'])} 条、覆盖 {fmt_num(values['partner_users'])} 人，样本规模不大但心理线索价值高，需要人工复核闭环")
    paragraph(doc, 17, f"漂流瓶匿名环境大幅降低表达门槛，回复率 {fmt_pct(values['bottle_reply_rate'], 1)} 证明同伴回应意愿强")
    paragraph(doc, 18, f"AI 模拟面试已服务 {fmt_num(values['interview_users'])} 名学生，报告转化率 {fmt_pct(values['interview_depth'], 1)}，可作为升学备考的集中训练工具")
    paragraph(doc, 19, "平台后续增长不应只依赖自然使用，应围绕高频问题、心理线索和升学节点设计运营动作")

    paragraph(doc, 26, "平台在本统计周期内呈现从集中启动、多模块放量到日常互动沉淀的阶段变化。相比单纯观察后台记录数，更重要的是识别每个阶段学生为什么使用平台、学校可以承接什么服务。")
    trend_parts = trend_analysis(registry)
    for idx, (title, text) in zip([30, 33, 36], trend_parts[:3]):
        paragraph(doc, idx, title)
        paragraph(doc, idx + 1, text)
    paragraph(doc, 39, trend_parts[3][1])

    paragraph(doc, 45, f"AI生涯规划模块聚焦生涯方向、专业选择与职业认知，是翼站最核心的功能之一。本周期覆盖 {fmt_num(values['career_users'])} 名学生，累计使用 {fmt_num(values['career_uses'])} 次，人均约 {values['career_avg']:.1f} 次，说明学生并非简单尝试，而是在围绕具体选择问题反复追问。")
    paragraph(doc, 51, insight_for_career(registry))
    career_insights = insight_for_career_table(registry)
    paragraph(doc, 52, career_insights[0])
    paragraph(doc, 53, career_insights[1])
    paragraph(doc, 58, "把学生分散的专业、学业和志愿疑问前置到平台端收集，帮助年级组形成可复用的生涯辅导议题库")
    paragraph(doc, 61, "关键判断：该功能当前最适合承接高二、高三的方向选择与升学决策需求，同时应把高频问题前置到高一生涯启蒙中，减少学生到关键节点才集中求助。")

    paragraph(doc, 67, f"★ AI 心语伙伴不是简单的聊天工具，而是学生主动表达心理困扰和情绪状态的低门槛入口。本周期覆盖 {fmt_num(values['partner_users'])} 人，说明使用面还需扩大；但已互动样本中的分级和困扰因子对学校具有线索价值。")
    partner_insights = insight_for_partner(registry)
    paragraph(doc, 74, partner_insights[2] if len(partner_insights) > 2 else partner_insights[0])
    paragraph(doc, 75, partner_insights[0])
    paragraph(doc, 76, partner_insights[1] if len(partner_insights) > 1 else partner_insights[0])
    state_insights = insight_for_partner_state(registry)
    paragraph(doc, 81, state_insights[2] if len(state_insights) > 2 else state_insights[0])
    paragraph(doc, 82, state_insights[0])
    paragraph(doc, 83, state_insights[1] if len(state_insights) > 1 else "心理状态分级的价值不在于替代教师判断，而在于把原本难以及时发现的学生表达转化为可复核线索，帮助学校把有限心理资源投向更需要关注的学生。")
    paragraph(doc, 86, "为学生提供低门槛表达入口，同时为学校沉淀可复核的心理支持线索")
    paragraph(doc, 89, "关键判断：该功能应纳入学校心理健康常态化工作，建议建立“AI 线索分级、班主任观察、心理教师复核”的闭环，把平台发现转化为可落地的支持名单和跟进流程。")

    paragraph(doc, 94, f"★ 漂流瓶是本周期最突出的互动场景，累计收发 {fmt_num(values['bottle_total'])} 条，回复率 {fmt_pct(values['bottle_reply_rate'], 1)}，说明学生不仅愿意表达，也愿意回应他人。")
    paragraph(doc, 98, f"平均 1 条发送可获得约 {(values['bottle_replies'] / values['bottle_posts']):.1f} 条回复，互动质量显著；但覆盖学生为发瓶学生去重，不包含所有回复参与者。")
    paragraph(doc, 105, "关键判断：该功能是低成本、高成效的校园文化建设抓手。学校应保留匿名表达的低门槛优势，同时通过正向话题和内容巡检把互动引导到支持性、建设性方向。")

    paragraph(doc, 107, "（四）AI 模拟面试：升学面试训练与反馈复盘工具")
    paragraph(doc, 109, f"AI 模拟面试已服务 {fmt_num(values['interview_users'])} 名学生，累计使用 {fmt_num(values['interview_records'])} 人次，生成报告 {fmt_num(values['interview_reports'])} 份。当前报告转化率为 {fmt_pct(values['interview_depth'], 1)}，说明已有一部分学生进入复盘环节，但仍有提升空间。")
    paragraph(doc, 112, "从使用人次与报告转化看，模拟面试已形成“练习—反馈—复盘”的初步闭环，后续重点应从学生自行试用转向学校有组织的训练安排。")
    paragraph(doc, 113, "建议围绕系统已开放的模拟练习、智能评分、语音回听和改进建议，形成集中练习、报告解读、改进清单和二次演练的流程。")
    paragraph(doc, 116, "形成标准化训练场景，便于教师掌握学生表达能力、答题结构和复盘完成情况")
    paragraph(doc, 117, "把学生分散练习转化为可追踪的报告材料，为后续辅导提供依据")
    paragraph(doc, 118, "通过集中讲评和二次演练提升面试训练的组织化水平")
    paragraph(doc, 120, "关键判断：该功能已形成一定报告转化，建议纳入学校升学服务的训练流程，由年级组或升学指导团队统一组织练习、讲评和复盘。")

    paragraph(doc, 126, f"漂流瓶 {fmt_pct(values['bottle_reply_rate'], 1)} 的回复率是平台高粘性的核心标志；AI 生涯规划人均 {values['career_avg']:.1f} 次使用说明学生愿意围绕复杂问题持续探索；AI 心语伙伴人均 {values['partner_avg']:.1f} 条对话则提示心理陪伴功能仍处在扩面阶段。整体看，平台已具备真实使用基础，但不同模块的成熟度不同，后续应分模块制定运营目标。")

    paragraph(doc, 128, "五、平台对学校的综合价值")

    header = doc.sections[0].header.paragraphs[0]
    first_run_text(header, f"{school} 生涯翼站使用情况学期分析报告  |  {datetime.now(TZ).year} 年 {datetime.now(TZ).month} 月")


def render_word_report(template_path, registry_path, output_path):
    registry = load_json(registry_path)
    if not registry.get("school_name") or len(registry.get("core_metrics", [])) < 5 or len(registry.get("module_metrics", [])) < 8:
        raise RuntimeError("Registry is incomplete; refusing to render an empty Word report.")

    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"Word template not found: {template}")
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(template, output)
    doc = Document(str(output))
    remove_template_time_labels(doc)
    values = fill_tables(doc, registry)
    fill_paragraphs(doc, registry, values)
    remove_template_time_labels(doc)
    apply_table_pagination_rules(doc)
    doc.save(str(output))
    return output


def main():
    parser = argparse.ArgumentParser(description="Render 生涯翼站 Word report from registry JSON.")
    parser.add_argument("--registry", required=True, help="Registry JSON generated by build_registry.py.")
    parser.add_argument("--out", required=True, help="Output .docx path.")
    parser.add_argument("--template", default=str(default_template_path()), help="Word template path.")
    args = parser.parse_args()
    print(render_word_report(args.template, args.registry, args.out))


if __name__ == "__main__":
    main()
