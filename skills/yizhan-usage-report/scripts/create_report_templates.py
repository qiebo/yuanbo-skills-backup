from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "assets" / "templates"
WORD_TEMPLATE = TEMPLATE_DIR / "yizhan-report-template.docx"


ACCENT = RGBColor(22, 122, 116)
DARK = RGBColor(23, 33, 43)
MUTED = RGBColor(104, 116, 131)
BLUE = RGBColor(47, 103, 162)
GREEN = RGBColor(90, 143, 75)
GOLD = RGBColor(194, 139, 27)
CORAL = RGBColor(216, 100, 69)
WHITE = RGBColor(255, 255, 255)
FONT_NAME = "微软雅黑"
HEADER_FILL = "1F4E78"
KEY_COLORS = [ACCENT, GOLD, GREEN, BLUE, CORAL]
MODULE_COLORS = {
    "career": ACCENT,
    "partner": GREEN,
    "bottle": GOLD,
    "interview": CORAL,
}


def set_run_font(run, size=10.5, bold=False, color=DARK, italic=False):
    run.font.name = FONT_NAME
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for theme_attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme", "hint"):
        r_fonts.attrib.pop(qn(f"w:{theme_attr}"), None)
    for font_attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{font_attr}"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def set_cell_text(cell, text, bold=False, color=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=9.5, bold=bold, color=color or DARK)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def set_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                shade_cell(cell, HEADER_FILL)
                set_cell_text(cell, text, bold=True, color=WHITE)
            else:
                set_cell_text(cell, text, color=DARK)
    return table


def color_table_row(table, row_idx, colors):
    row = table.rows[row_idx]
    for idx, cell in enumerate(row.cells):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=11, bold=True, color=colors[idx % len(colors)])


def color_table_column(table, col_idx, color):
    for row in table.rows[1:]:
        cell = row.cells[col_idx]
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=10, bold=True, color=color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=18 if level == 1 else 13, bold=True, color=ACCENT if level == 1 else DARK)
    return p


def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, italic=italic, color=MUTED if italic else DARK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("第 ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    paragraph.add_run(" 页")


def create_word_template():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    styles["Normal"].font.name = FONT_NAME
    normal_fonts = styles["Normal"]._element.get_or_add_rPr().get_or_add_rFonts()
    for theme_attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme", "hint"):
        normal_fonts.attrib.pop(qn(f"w:{theme_attr}"), None)
    for font_attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal_fonts.set(qn(f"w:{font_attr}"), FONT_NAME)
    styles["Normal"].font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.text = "{{school_name}} 生涯翼站使用情况学期分析报告 | {{report_date}}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("{{school_name}}")
    set_run_font(run, size=28, bold=True, color=DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("生涯翼站使用情况学期分析报告")
    set_run_font(run, size=20, bold=True, color=ACCENT)

    meta = add_table(doc, [
        ["数据周期", "{{period_text}}"],
        ["报告日期", "{{report_date}}"],
        ["数据来源", "生涯翼站后台聚合数据 / 学校统计查询表"],
        ["报告口径", "本报告固定区分去重服务覆盖率、登录人次覆盖强度、后台记录数趋势"],
    ])
    for row in meta.rows:
        set_width(row.cells[0], 4)
        set_width(row.cells[1], 11)

    add_heading(doc, "一、执行摘要", 1)
    add_body(doc, "{{executive_summary}}")
    metrics_table = add_table(doc, [
        ["累计服务学生", "累计交互总量", "去重服务覆盖率", "登录人次覆盖强度", "活跃峰值"],
        ["{{served_students}}", "{{total_interactions}}", "{{service_coverage_rate}}", "{{login_intensity_rate}}", "{{active_peak}}"],
    ])
    color_table_row(metrics_table, 1, KEY_COLORS)
    add_heading(doc, "核心发现", 2)
    add_bullets(doc, [
        "{{finding_1}}",
        "{{finding_2}}",
        "{{finding_3}}",
        "{{finding_4}}",
    ])

    add_heading(doc, "二、趋势与模块贡献", 1)
    add_body(doc, "{{trend_summary}}")
    add_body(doc, "口径说明：本节趋势为网站后台记录数趋势，不等同于功能使用次数趋势。", italic=True)
    add_table(doc, [
        ["月份", "AI 生涯规划记录", "AI 心语伙伴记录", "漂流瓶发布记录", "AI 模拟面试报告", "后台记录合计"],
        ["{{trend_month_1}}", "{{career_records_1}}", "{{partner_records_1}}", "{{bottle_records_1}}", "{{interview_reports_1}}", "{{total_records_1}}"],
        ["{{trend_month_2}}", "{{career_records_2}}", "{{partner_records_2}}", "{{bottle_records_2}}", "{{interview_reports_2}}", "{{total_records_2}}"],
        ["{{trend_month_3}}", "{{career_records_3}}", "{{partner_records_3}}", "{{bottle_records_3}}", "{{interview_reports_3}}", "{{total_records_3}}"],
        ["{{trend_month_4}}", "{{career_records_4}}", "{{partner_records_4}}", "{{bottle_records_4}}", "{{interview_reports_4}}", "{{total_records_4}}"],
        ["{{trend_month_5}}", "{{career_records_5}}", "{{partner_records_5}}", "{{bottle_records_5}}", "{{interview_reports_5}}", "{{total_records_5}}"],
        ["{{trend_month_6}}", "{{career_records_6}}", "{{partner_records_6}}", "{{bottle_records_6}}", "{{interview_reports_6}}", "{{total_records_6}}"],
    ])

    add_heading(doc, "三、功能模块深度分析", 1)
    modules = [
        ("AI 生涯规划", "career", "新高考选科决策的科学支撑"),
        ("AI 心语伙伴", "partner", "学生心理支持与风险初筛补充"),
        ("漂流瓶", "bottle", "校园互助文化的自发载体"),
        ("AI 模拟面试", "interview", "升学面试训练与反馈复盘工具"),
    ]
    for name, key, subtitle_text in modules:
        add_heading(doc, f"（{len([m for m in modules if modules.index(m) <= modules.index((name, key, subtitle_text))])}）{name}：{subtitle_text}", 2)
        add_body(doc, "{{" + key + "_analysis}}")
        module_table = add_table(doc, [
            ["核心指标", "数值", "口径说明"],
            ["{{" + key + "_kpi_1_name}}", "{{" + key + "_kpi_1_value}}", "{{" + key + "_kpi_1_note}}"],
            ["{{" + key + "_kpi_2_name}}", "{{" + key + "_kpi_2_value}}", "{{" + key + "_kpi_2_note}}"],
            ["{{" + key + "_kpi_3_name}}", "{{" + key + "_kpi_3_value}}", "{{" + key + "_kpi_3_note}}"],
        ])
        color_table_column(module_table, 1, MODULE_COLORS[key])
        add_heading(doc, "学校行动建议", 2)
        add_body(doc, "{{" + key + "_action}}")

    add_heading(doc, "四、重点内容与年级画像", 1)
    add_table(doc, [
        ["模块", "画像维度", "关键发现", "学校可采取动作"],
        ["AI 生涯规划", "{{career_profile_dimension}}", "{{career_profile_finding}}", "{{career_profile_action}}"],
        ["AI 心语伙伴", "{{partner_profile_dimension}}", "{{partner_profile_finding}}", "{{partner_profile_action}}"],
        ["漂流瓶", "{{bottle_profile_dimension}}", "{{bottle_profile_finding}}", "{{bottle_profile_action}}"],
        ["AI 模拟面试", "{{interview_profile_dimension}}", "{{interview_profile_finding}}", "{{interview_profile_action}}"],
    ])

    add_heading(doc, "五、学校行动建议与下阶段运营重点", 1)
    add_table(doc, [
        ["建议方向", "具体建议", "优先级"],
        ["分年级运营", "{{grade_operation_advice}}", "{{grade_operation_priority}}"],
        ["心理支持闭环", "{{mental_support_advice}}", "{{mental_support_priority}}"],
        ["数据复盘机制", "{{data_review_advice}}", "{{data_review_priority}}"],
    ])

    add_heading(doc, "六、数据口径与限制说明", 1)
    add_bullets(doc, [
        "去重服务覆盖率 = 累计服务学生 / 学生账号总数。",
        "登录人次覆盖强度 = 总登录人次 / 学生账号总数，不能等同于去重学生覆盖率。",
        "网站趋势为后台记录数趋势，不等同于使用次数趋势。",
        "漂流瓶覆盖学生如使用 sender_id 去重，只代表发瓶学生，不包含回复参与者。",
        "AI 心语伙伴历史分析口径与内容统计口径必须分开解释。",
        "{{additional_limitation_note}}",
    ])

    doc.save(WORD_TEMPLATE)


if __name__ == "__main__":
    TEMPLATE_DIR.mkdir(exist_ok=True)
    create_word_template()
    print(WORD_TEMPLATE)
