#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
远播方案专家团 · 环境就绪自检（Pre-flight Environment Check）

用途
----
每次启动方案对话的第一动作由「方案总师」运行本脚本，提前把排版所需的环境
依赖（当前唯一外部依赖：python-docx）补齐，避免流程跑到 Phase 7.5「精美
Word 精排版」时才临时装包导致的中途卡顿。

设计原则
--------
1. 纯标准库，不依赖任何第三方包，保证本脚本自身永远可运行。
2. 检测「运行本脚本的解释器」是否可 import python-docx。
3. 默认**只检测、不安装**：缺失时输出 DOCX_READY=no 与缺失清单，由总师
   经 AskUserQuestion 征得用户授权后，带 `--install` 重新运行本脚本才执行
   安装（避免在受限网络/企业 IT 环境下擅自安装软件）。
4. `--install` 时按优先顺序自动补齐：
   a) 直接 `python -m pip install python-docx`（装到当前解释器）；
   b) 失败且为 externally-managed（Linux 系统 Python 常见）→ 在隔离
      venv（~/.workbuddy/skill_env_cache/...）里安装，并回报该 venv 解释器；
   c) 仍失败 → `pip install --user`；
   d) 全部失败 → 退出码 2，给出人工处置建议，不静默放过。
5. 装好后做最小功能回环（建一个临时 docx 再读回），确认真能用，而非仅 import 通过。

退出码
------
0  python-docx 就绪（已具备或已补齐）
1  未就绪且未授权安装（默认只检测；加 --install 才尝试安装）
2  致命：无任何可用解释器 / 安装失败，需人工介入

用法
----
  检查（默认，不安装）：python bin/check_env.py
  授权安装：             python bin/check_env.py --install

机器可读输出（供 Agent 解析）
----------------------------
DOCX_PYTHON=<绝对路径>    后续 Phase 7.5 排版应使用此解释器
DOCX_READY=yes|no
"""

from __future__ import annotations

import os
import subprocess
import sys
import tempfile
from pathlib import Path


REQUIRED_PKG = "python-docx"
VENV_REL = os.path.join(
    ".workbuddy", "skill_env_cache", "yuanbo-school-proposal-team", "venv"
)


def configure_stdio() -> None:
    """Keep the preflight report readable on Windows GBK consoles."""
    for stream in (sys.stdout, sys.stderr):
        reconfigure = getattr(stream, "reconfigure", None)
        if reconfigure:
            reconfigure(encoding="utf-8", errors="replace")


def banner(msg: str) -> None:
    print("=" * 56)
    print(msg)
    print("=" * 56)


def run(cmd):
    """运行子进程，返回 (returncode, stdout, stderr)。"""
    try:
        proc = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120,
        )
        return proc.returncode, proc.stdout, proc.stderr
    except Exception as exc:  # noqa: BLE001
        return 1, "", f"subprocess error: {exc}"


def can_import(target_py: str) -> bool:
    rc, _, _ = run(
        [target_py, "-c", "import docx; from docx import Document; _ = docx.__version__"]
    )
    return rc == 0


def functional_check(target_py: str) -> bool:
    """最小功能回环：建临时 docx 并读回，确认真能产出 .docx。"""
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "probe.docx"
        rc, _, err = run(
            [
                target_py,
                "-c",
                (
                    "from docx import Document;"
                    f"d=Document();d.add_heading('probe',0);"
                    f"d.add_paragraph('env_ok');d.save({str(out)!r});"
                ),
            ]
        )
        if rc != 0:
            print(f"  [功能回环失败] {err.strip()[:200]}")
            return False
        if not out.exists():
            print("  [功能回环失败] 未生成 docx 文件")
            return False
        return True


def try_install(target_py: str, extra_args=None) -> tuple[bool, str]:
    cmd = [target_py, "-m", "pip", "install", REQUIRED_PKG]
    if extra_args:
        cmd += extra_args
    print(f"  [安装] {' '.join(cmd)}")
    rc, out, err = run(cmd)
    if rc == 0:
        return True, ""
    return False, (out + err).strip()


def ensure_via_venv() -> tuple[bool, str]:
    """externally-managed 场景：在隔离 venv 中安装，回报 venv 解释器。"""
    home = Path.home()
    venv_dir = home / VENV_REL
    venv_py = (
        venv_dir / "Scripts" / "python.exe"
        if os.name == "nt"
        else venv_dir / "bin" / "python"
    )
    if not venv_py.exists():
        print(f"  [venv] 创建隔离环境 {venv_dir}")
        rc, out, err = run([sys.executable, "-m", "venv", str(venv_dir)])
        if rc != 0:
            return False, f"venv 创建失败: {err.strip()[:200]}"
    ok, detail = try_install(str(venv_py), None)
    if ok:
        return True, str(venv_py)
    return False, detail


def main() -> int:
    import argparse
    configure_stdio()
    ap = argparse.ArgumentParser(description="远播方案专家团环境就绪自检")
    ap.add_argument("--install", action="store_true", help="授权自动安装缺失依赖（默认只检测不安装）")
    args = ap.parse_args()

    banner("远播方案专家团 · 环境就绪自检")
    target = sys.executable
    print(f"解释器：{target}")
    print(f"依赖项：{REQUIRED_PKG}")

    # 1) 已就绪
    if can_import(target):
        print("状态：python-docx 已具备 ✅")
        if not functional_check(target):
            print("状态：import 通过但功能回环失败，尝试重装…")
        else:
            print("DOCX_PYTHON=" + target)
            print("DOCX_READY=yes")
            return 0

    # 依赖缺失：默认只报告、不自动安装（企业 IT 环境不擅自装软件）
    if not args.install:
        print("状态：python-docx 缺失 ❌")
        print("缺失依赖：python-docx")
        print("已停止自动安装（默认只检测）。征得用户同意后，任选其一：")
        print(f"  a) {target} -m pip install python-docx")
        print("  b) 重新运行本脚本并加 --install 参数（脚本自动补齐）")
        print("DOCX_PYTHON=" + target)
        print("DOCX_READY=no")
        return 1

    print("状态：python-docx 缺失，用户已授权，开始自动补齐…")

    # 2) 直接装到当前解释器
    ok, detail = try_install(target, None)
    if ok and can_import(target):
        print("状态：已就地安装并导入成功 ✅")
        if not functional_check(target):
            print("  [警告] 功能回环未通过，请检查安装。")
            print("DOCX_PYTHON=" + target)
            print("DOCX_READY=no")
            return 2
        print("DOCX_PYTHON=" + target)
        print("DOCX_READY=yes")
        return 0

    # 3) externally-managed → venv
    if "externally-managed-environment" in detail:
        print("  [检测到 externally-managed] 改用隔离 venv")
        ok, venv_py = ensure_via_venv()
        if ok and can_import(venv_py):
            print(f"状态：已在隔离 venv 安装成功 ✅（{venv_py}）")
            print("DOCX_PYTHON=" + venv_py)
            print("DOCX_READY=yes")
            return 0
        detail = venv_py

    # 4) --user 兜底
    ok, detail2 = try_install(target, ["--user"])
    if ok and can_import(target):
        print("状态：已通过 --user 安装成功 ✅")
        print("DOCX_PYTHON=" + target)
        print("DOCX_READY=yes")
        return 0

    # 5) 全部失败
    banner("❌ 环境补齐失败（致命）")
    print("无法自动安装 python-docx，可能原因：无网络、pip 受限、权限不足。")
    print("人工处置（任选其一）：")
    print(f"  a) {target} -m pip install python-docx")
    print(f"  b) {target} -m pip install --user python-docx")
    print("  c) 安装 WorkBuddy 托管 Python 后重试（推荐，隔离不污染系统）")
    if detail2:
        print("最后错误：")
        print("  " + detail2[-500:].replace("\n", "\n  "))
    print("DOCX_PYTHON=" + target)
    print("DOCX_READY=no")
    return 2


if __name__ == "__main__":
    sys.exit(main())
