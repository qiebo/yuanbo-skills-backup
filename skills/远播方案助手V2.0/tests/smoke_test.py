#!/usr/bin/env python3
"""端到端冒烟测试：校验、构建并检查表格未丢失。"""
from __future__ import annotations

import json
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


def run(cmd):
    proc = subprocess.run([str(x) for x in cmd], text=True, capture_output=True)
    if proc.returncode != 0:
        raise RuntimeError(f"命令失败：{' '.join(map(str, cmd))}\n{proc.stdout}\n{proc.stderr}")
    return proc


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    sample = root / "tests/sample_project"
    with tempfile.TemporaryDirectory(prefix="yb_smoke_") as td:
        project = Path(td) / "sample_project"
        shutil.copytree(sample, project, ignore=shutil.ignore_patterns("output", "qa", "__pycache__"))
        run([sys.executable, root / "scripts/validate_project.py", project])
        run([sys.executable, root / "scripts/build_docx.py", project])
        manifest = json.loads((project / "qa/build_manifest.json").read_text(encoding="utf-8"))
        doc = Document(manifest["output"])
        assert len(doc.tables) == 2, f"预期2个表格，实际{len(doc.tables)}"
        assert any("项目背景与目标" in p.text for p in doc.paragraphs), "正文标题缺失"
        assert any("目录" in p.text for p in doc.paragraphs), "静态目录缺失"
    print("smoke test passed")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
