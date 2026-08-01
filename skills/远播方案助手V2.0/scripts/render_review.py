#!/usr/bin/env python3
"""将 DOCX 渲染为 PDF 与逐页 PNG，供视觉检查。"""
from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import tempfile
from pathlib import Path

try:
    from docx import Document as _DocxDocument
except Exception:  # pragma: no cover
    _DocxDocument = None


def require(name: str) -> str:
    path = shutil.which(name)
    if not path:
        raise SystemExit(f"缺少工具：{name}")
    return path


def structural_check(docx: Path, out: Path) -> int:
    """LibreOffice 不可用时的结构化降级验收（非视觉渲染）。

    用 python-docx 做结构层检查（段落/表格/图片计数与启发式），并明确标注
    "不能替代视觉验收"，提示人工在 Word 中复核版式。
    """
    issues = []
    paras = []
    tables = []
    images = 0
    if _DocxDocument is None:
        issues.append("缺少 python-docx，结构降级检查也无法执行，请安装后重试。")
    else:
        doc = _DocxDocument(str(docx))
        paras = [p for p in doc.paragraphs if p.text.strip()]
        tables = doc.tables
        images = len(doc.inline_shapes)
        if not paras:
            issues.append("正文为空：没有非空段落。")
        for ti, t in enumerate(tables):
            try:
                cols = len(t.columns)
            except Exception:
                cols = 0
            if cols > 6:
                issues.append(f"表格 {ti + 1} 列数较多（{cols} 列），A4 纵向可能越界，请人工核对（必要时改为横向或拆分）。")
        if images == 0:
            issues.append("未检测到内嵌图片；若方案应含效果图/图纸，请核对图片是否成功嵌入。")
        for ti, t in enumerate(tables):
            for row in t.rows:
                if any(not c.text.strip() for c in row.cells) and row.cells:
                    # 仅提示空单元格，可能是合并单元格，不误报为错误
                    pass

    checklist = out / "结构降级检查清单.md"
    lines = [
        "# 结构降级检查清单（非视觉渲染）",
        "",
        "> 本机未安装 LibreOffice / soffice，无法进行逐页 PNG 视觉渲染。",
        "> 以下为基于 python-docx 的结构化检查，**不能替代视觉验收**，请人工在 Word 中复核版式后再交付。",
        "",
        f"- DOCX：`{docx.name}`",
        f"- 非空段落数：{len(paras)}",
        f"- 表格数：{len(tables)}",
        f"- 内嵌图片数：{images}",
        "",
        "## 结构项（人工勾选）",
        "",
        "- [ ] 封面独立、目录存在且不为空",
        "- [ ] 标题层级正确（一级/二级/三级）",
        "- [ ] 表格列数与页面方向匹配（无越界）",
        "- [ ] 图片存在、清晰且无占位符",
        "- [ ] 无“示例 / 待确认 / TODO / 占位”残留",
        "- [ ] 页眉页脚、页码、字体、段首缩进一致",
        "",
        "## 自动发现问题",
        "",
    ]
    if issues:
        lines += [f"- ⚠ {i}" for i in issues]
    else:
        lines += ["结构层未发现明显问题（仍需人工视觉复核）。"]
    checklist.write_text("\n".join(lines) + "\n", encoding="utf-8")

    report = {
        "docx": str(docx),
        "mode": "structural_fallback",
        "visual_rendered": False,
        "paragraphs": len(paras),
        "tables": len(tables),
        "images": images,
        "issues": issues,
        "checklist": str(checklist),
    }
    (out / "render_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print("未检测到 LibreOffice / soffice：已执行结构降级验收（非视觉渲染），请人工复核版式。")
    print(checklist)
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="渲染 DOCX 进行视觉验收")
    parser.add_argument("docx")
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--dpi", type=int, default=150)
    args = parser.parse_args()

    docx = Path(args.docx).expanduser().resolve()
    out = Path(args.output_dir).expanduser().resolve()
    if not docx.exists():
        raise SystemExit(f"文件不存在：{docx}")
    out.mkdir(parents=True, exist_ok=True)
    for stale in list(out.glob("page-*.png")) + list(out.glob("*.pdf")):
        stale.unlink()
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not libreoffice:
        # 降级：LibreOffice 不可用时执行结构化验收（非视觉渲染）
        return structural_check(docx, out)
    pdftoppm = require("pdftoppm")

    with tempfile.TemporaryDirectory(prefix="yb_lo_") as profile:
        env = os.environ.copy()
        env["HOME"] = profile
        cmd = [libreoffice, "--headless", f"-env:UserInstallation=file://{profile}/profile", "--convert-to", "pdf", "--outdir", str(out), str(docx)]
        proc = subprocess.run(cmd, text=True, capture_output=True, env=env)
        pdf = out / f"{docx.stem}.pdf"
        if proc.returncode != 0 or not pdf.exists():
            raise SystemExit(f"LibreOffice 渲染失败：\nSTDOUT: {proc.stdout}\nSTDERR: {proc.stderr}")

    prefix = out / "page"
    proc = subprocess.run([pdftoppm, "-png", "-r", str(args.dpi), str(pdf), str(prefix)], text=True, capture_output=True)
    if proc.returncode != 0:
        raise SystemExit(f"PDF 转 PNG 失败：{proc.stderr}")

    pages = sorted(out.glob("page-*.png"))
    if not pages:
        raise SystemExit("没有生成页面图片。")
    checklist = out / "视觉检查清单.md"
    lines = [
        "# 视觉检查清单", "", f"- DOCX：`{docx.name}`", f"- 页面数：{len(pages)}", "",
        "逐页以 100% 缩放检查：", "",
        "- [ ] 封面独立且无不应出现的页眉页脚", "- [ ] 目录存在且不为空", "- [ ] 正文标题层级和分页正确",
        "- [ ] 表格无越界、缺列、断裂或小字不可读", "- [ ] 图片无变形、遮挡和低清晰度", "- [ ] 无空白页、占位符或未替换标记",
        "- [ ] 页眉页脚、页码、字体和行距一致", "", "## 页面", ""
    ]
    lines.extend(f"- [ ] {p.name}" for p in pages)
    checklist.write_text("\n".join(lines) + "\n", encoding="utf-8")
    report = {"docx": str(docx), "pdf": str(pdf), "pages": [str(p) for p in pages], "page_count": len(pages), "checklist": str(checklist), "mode": "visual", "visual_rendered": True}
    (out / "render_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"渲染完成：{len(pages)} 页")
    print(checklist)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
