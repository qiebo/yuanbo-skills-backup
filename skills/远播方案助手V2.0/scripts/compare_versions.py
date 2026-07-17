#!/usr/bin/env python3
"""比较两个 Markdown 或 DOCX 版本的文本差异。"""
from __future__ import annotations

import argparse
import difflib
from pathlib import Path

from docx import Document


def extract(path: Path) -> list[str]:
    if path.suffix.lower() == ".docx":
        doc = Document(path)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for ti, table in enumerate(doc.tables, 1):
            lines.append(f"[表格{ti}]")
            lines.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return lines
    return path.read_text(encoding="utf-8").splitlines()


def main() -> int:
    parser = argparse.ArgumentParser(description="比较方案版本")
    parser.add_argument("old")
    parser.add_argument("new")
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    old, new = Path(args.old), Path(args.new)
    diff = difflib.unified_diff(extract(old), extract(new), fromfile=old.name, tofile=new.name, lineterm="")
    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(diff) + "\n", encoding="utf-8")
    print(f"差异已写入：{out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
