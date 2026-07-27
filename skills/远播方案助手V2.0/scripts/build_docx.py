#!/usr/bin/env python3
"""从项目终稿构建正式 Word。Pandoc 负责 Markdown 解析，python-docx 负责封面和版式后处理。"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def require_tool(name: str) -> str:
    path = shutil.which(name)
    if not path:
        raise SystemExit(f"缺少构建工具 {name}。请安装后重试；禁止回退到示例内容。")
    return path


def clear_container(container) -> None:
    for p in list(container.paragraphs):
        p._element.getparent().remove(p._element)


def add_page_field(paragraph, start_at_zero: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    run.font.size = Pt(10)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    paragraph.add_run(" 页")


def set_page_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant = tr_pr.find(qn("w:cantSplit"))
    if cant is None:
        cant = OxmlElement("w:cantSplit")
        tr_pr.append(cant)


def insert_front_matter(doc: Document, project: dict, output_cfg: dict, markdown_text: str) -> None:
    """插入封面和确定性静态目录，位于 Pandoc 正文之前。"""
    style_names = {s.name for s in doc.styles}
    created = []

    spacer = doc.add_paragraph("")
    spacer.paragraph_format.space_before = Pt(120)
    created.append(spacer._p)

    title = doc.add_paragraph(project.get("title", project.get("name", "项目方案")))
    title.style = "Cover Title" if "Cover Title" in style_names else "Title"
    created.append(title._p)

    client = doc.add_paragraph(project.get("client", ""))
    client.style = "Cover Client" if "Cover Client" in style_names else "Subtitle"
    created.append(client._p)

    version = doc.add_paragraph(f"版本：{project.get('version', 'v1.0')}")
    version.style = "Cover Meta" if "Cover Meta" in style_names else "Subtitle"
    created.append(version._p)

    date = doc.add_paragraph(project.get("date", ""))
    date.style = "Cover Meta" if "Cover Meta" in style_names else "Subtitle"
    created.append(date._p)

    if project.get("confidentiality"):
        conf = doc.add_paragraph(project["confidentiality"])
        conf.style = "Cover Meta" if "Cover Meta" in style_names else "Subtitle"
        created.append(conf._p)

    cover_break = doc.add_paragraph("")
    cover_break.add_run().add_break(WD_BREAK.PAGE)
    created.append(cover_break._p)

    toc_mode = output_cfg.get("toc_mode", "static" if output_cfg.get("toc", True) else "none")
    if toc_mode == "static":
        toc_title = doc.add_paragraph("目录")
        toc_title.style = "TOC Heading" if "TOC Heading" in style_names else "Heading 1"
        created.append(toc_title._p)
        headings = []
        for line in markdown_text.splitlines():
            match = re.match(r"^(#{1,3})\s+(.+?)\s*$", line)
            if match:
                headings.append((len(match.group(1)), re.sub(r"[*_`]+", "", match.group(2))))
        for level, text in headings:
            entry = doc.add_paragraph(text)
            entry.paragraph_format.left_indent = Cm(0.65 * (level - 1))
            entry.paragraph_format.space_before = Pt(2)
            entry.paragraph_format.space_after = Pt(4)
            entry.paragraph_format.keep_with_next = False
            if entry.runs:
                entry.runs[0].font.size = Pt(12 if level == 1 else 11)
                entry.runs[0].bold = level == 1
            created.append(entry._p)
        toc_break = doc.add_paragraph("")
        toc_break.add_run().add_break(WD_BREAK.PAGE)
        created.append(toc_break._p)

    body = doc._element.body
    for element in reversed(created):
        body.insert(0, element)



def set_cell_border(cell, color: str = "A6A6A6", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


DEFAULT_FONT = "微软雅黑"
# 首行缩进单位按"字符数"理解（中文排版惯例：正文段首空两格 = 2 字符）
DEFAULT_FIRST_LINE_INDENT_CHARS = 2


def set_run_font(run, font: str) -> None:
    """同时设置 eastAsia/ascii/hAnsi，确保中文与西文都落到指定字体。"""
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)


def apply_font_and_indent(doc: Document, font: str, indent_chars) -> None:
    """统一字体（正文/表格单元格/页眉页脚/样式默认）+ 正文段首缩进（按字符数）。

    - 字体应用到所有 run（含表格单元格、页眉页脚），并写入样式默认值兜底未显式 run。
    - 首行缩进仅作用于正文叙事段，跳过标题/封面/目录/空段；表格单元格显式不缩进。
    """
    # 1) 样式默认字体兜底
    for style in doc.styles:
        try:
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), font)
            rfonts.set(qn("w:ascii"), font)
            rfonts.set(qn("w:hAnsi"), font)
        except Exception:
            pass
    # 2) 正文段落 run 字体
    for p in doc.paragraphs:
        for run in p.runs:
            set_run_font(run, font)
    # 3) 表格单元格 run 字体 + 显式不缩进
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.left_indent = Pt(0)
                    for run in p.runs:
                        set_run_font(run, font)
    # 4) 页眉页脚 run 字体
    for section in doc.sections:
        for hf in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            for p in hf.paragraphs:
                for run in p.runs:
                    set_run_font(run, font)
    # 5) 正文段首缩进（跳过标题/封面/目录/空段）
    try:
        chars = float(indent_chars)
    except (TypeError, ValueError):
        chars = DEFAULT_FIRST_LINE_INDENT_CHARS if indent_chars else 0
    if chars and chars > 0:
        for p in doc.paragraphs:
            style_name = p.style.name if p.style else ""
            if any(k in style_name for k in ("Heading", "Cover", "Title", "TOC")):
                continue
            if p.paragraph_format.left_indent is not None:  # 目录条目等带左缩进的段落
                continue
            if not p.text.strip():
                continue
            size = None
            for run in p.runs:
                if run.font.size:
                    size = run.font.size.pt
                    break
            if size is None:
                size = 12
            p.paragraph_format.first_line_indent = Pt(chars * size)


def postprocess(docx_path: Path, project: dict, output_cfg: dict, markdown_text: str) -> dict:
    doc = Document(docx_path)
    insert_front_matter(doc, project, output_cfg, markdown_text)
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.5)

    clear_container(section.first_page_header)
    clear_container(section.first_page_footer)
    clear_container(section.header)
    header = section.header.add_paragraph(output_cfg.get("header_text", project.get("title", "")))
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.size = Pt(9)
    clear_container(section.footer)
    if output_cfg.get("footer_page_number", True):
        add_page_field(section.footer.add_paragraph())
        set_page_start(section, 0)

    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True

    style_names = {s.name for s in doc.styles}
    for table in doc.tables:
        table.style = "Table" if "Table" in style_names else ("Table Grid" if "Table Grid" in style_names else None)
        if table.rows:
            set_repeat_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            set_cant_split(row)
            for cell_index, cell in enumerate(row.cells):
                set_cell_border(cell)
                if row_index == 0:
                    set_cell_shading(cell)
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.first_line_indent = Pt(0)
                    if row_index == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9.5)
                        if row_index == 0:
                            run.bold = True

    # 统一字体与正文段首缩进（可配置 output.font / output.first_line_indent）
    apply_font_and_indent(
        doc,
        output_cfg.get("font", DEFAULT_FONT),
        output_cfg.get("first_line_indent", DEFAULT_FIRST_LINE_INDENT_CHARS),
    )

    set_update_fields(doc)
    doc.core_properties.title = project.get("title", "")
    doc.core_properties.subject = project.get("name", "")
    doc.core_properties.author = "远播方案助手V2.0"
    doc.core_properties.comments = "由结构化源文件生成；正文修改应回到 Markdown 与数据台账。"
    doc.save(docx_path)

    markdown_tables = len(re.findall(r"(?m)^\s*\|.+\|\s*$", markdown_text))
    output_tables = len(doc.tables)
    if markdown_tables and output_tables == 0:
        raise SystemExit("检测到 Markdown 表格，但 Word 中没有表格。已停止交付。")
    return {"paragraphs": len(doc.paragraphs), "tables": output_tables, "sections": len(doc.sections)}


def main() -> int:
    parser = argparse.ArgumentParser(description="构建方案 Word")
    parser.add_argument("project_dir")
    parser.add_argument("--force", action="store_true", help="允许在校验未通过时构建，仅用于调试")
    parser.add_argument("--overwrite", action="store_true", help="覆盖写同名输出文件；若已存在先备份为 _备份_时间戳，避免时间戳改名导致的后处理错位")
    args = parser.parse_args()

    pandoc = require_tool("pandoc")
    project_dir = Path(args.project_dir).expanduser().resolve()
    config_path = project_dir / "project.yaml"
    source = project_dir / "07_方案终稿.md"
    if not config_path.exists() or not source.exists():
        raise SystemExit("缺少 project.yaml 或 07_方案终稿.md；禁止使用示例内容回退。")

    validation = project_dir / "qa/validation_report.json"
    if validation.exists() and not args.force:
        report = json.loads(validation.read_text(encoding="utf-8"))
        if not report.get("passed"):
            raise SystemExit("确定性校验未通过。请先修复 qa/validation_report.md，或仅调试时使用 --force。")

    config = yaml.safe_load(config_path.read_text(encoding="utf-8")) or {}
    project = config.get("project", {})
    output_cfg = config.get("output", {})
    skill_root = Path(__file__).resolve().parents[1]
    stylepack = skill_root / "stylepacks" / output_cfg.get("stylepack", "学校建设方案版.docx")
    if not stylepack.exists():
        raise SystemExit(f"找不到样式包：{stylepack}")

    output_dir = project_dir / "output"
    output_dir.mkdir(exist_ok=True)
    filename = output_cfg.get("filename") or f"{project.get('name', '方案')}_{project.get('version', 'v1.0')}.docx"
    output_path = output_dir / filename
    backup_path = None
    if output_path.exists():
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if args.overwrite:
            # 覆盖模式：先备份旧版，再以干净文件名重写，避免后续后处理错位
            backup_path = output_path.with_name(f"{output_path.stem}_备份_{stamp}{output_path.suffix}")
            shutil.copy2(output_path, backup_path)
        else:
            output_path = output_path.with_name(f"{output_path.stem}_{stamp}{output_path.suffix}")

    markdown_text = source.read_text(encoding="utf-8")
    with tempfile.TemporaryDirectory(prefix="yb_docx_") as td:
        temp_docx = Path(td) / "body.docx"
        cmd = [
            pandoc, str(source),
            "--from=markdown+pipe_tables+footnotes+task_lists+strikeout",
            "--to=docx", "--standalone",
            f"--reference-doc={stylepack}",
            f"--resource-path={project_dir}:{project_dir / 'assets'}:{project_dir / 'sources'}",
            "--metadata=lang:zh-CN", "--metadata=toc-title:目录",
            "--output", str(temp_docx),
        ]
        toc_mode = output_cfg.get("toc_mode", "static" if output_cfg.get("toc", True) else "none")
        if toc_mode == "field":
            cmd.append("--toc")
        if output_cfg.get("numbered_sections", False):
            cmd.append("--number-sections")
        proc = subprocess.run(cmd, cwd=project_dir, text=True, capture_output=True)
        if proc.returncode != 0:
            raise SystemExit(f"Pandoc 构建失败：\n{proc.stderr}")
        shutil.copy2(temp_docx, output_path)

    stats = postprocess(output_path, project, output_cfg, markdown_text)
    manifest = {
        "source": str(source), "output": str(output_path), "stylepack": str(stylepack),
        "project_version": project.get("version"), "stats": stats,
        "backup": str(backup_path) if backup_path else None,
        "font": output_cfg.get("font", DEFAULT_FONT),
        "first_line_indent": output_cfg.get("first_line_indent", DEFAULT_FIRST_LINE_INDENT_CHARS),
        "built_at": datetime.now().isoformat(timespec="seconds")
    }
    (project_dir / "qa/build_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Word 已生成：{output_path}")
    if backup_path:
        print(f"旧版已备份：{backup_path}")
    print(f"段落 {stats['paragraphs']}，表格 {stats['tables']}，节 {stats['sections']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
