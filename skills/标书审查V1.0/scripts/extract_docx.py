#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Extract a .docx into structured Markdown for bid/tender review (V1.1).
Usage:
    python extract_docx.py <input.docx> [output.md] [--extract-images] [--image-dir output/images]
Design notes (built from real bid-review experience):
- Preserves Heading / Title styles as Markdown '#'..'######' so the reviewer
  can navigate the document structure.
- Tables are emitted as Markdown tables (merged cells via gridSpan/vMerge expanded).
- Paragraphs / cells / headers / footers that contain an image (certificates,
  screenshots, seals, contracts, PPT, VML pictures) are exported to
  <image-dir>/<safe>/img_NNN.png (where <safe> = <stem>_<md5-of-source>[:6]) and
  flagged with `[IMAGE: images/<safe>/img_NNN.png]` so the evidence is verifiable
  against the original file (see 铁律 2). The hash in <safe> keeps images of two
  different files that share a base name in separate subdirs.
- Position markers `<!-- §N ¶约第P页 -->` are emitted before each block to let
  reviewers cite evidence locations (page is an estimate from w:br page breaks).
"""
import sys
import os
import argparse
import hashlib
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
# Namespace URIs (kept explicit for VML which python-docx's qn may not map).
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
VML = 'urn:schemas-microsoft-com:vml'
O = 'urn:schemas-microsoft-com:office:office'
NS = {'w': W, 'r': R, 'a': A, 'v': VML, 'o': O}
def _tag(prefix, local):
    return '{%s}%s' % (NS[prefix], local)
def iter_block_items(parent):
    """Yield (kind, item) for paragraphs and tables in document order."""
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            yield ('p', Paragraph(child, parent))
        elif child.tag == qn('w:tbl'):
            yield ('t', Table(child, parent))
def style_level(p):
    """Return heading level (1-6) for heading/title styles, else None."""
    try:
        s = p.style.name
    except Exception:
        return None
    if not s:
        return None
    if s == 'Title' or s == '标题':
        return 1
    if 'Heading' in s or '标题' in s:
        digits = ''.join(ch for ch in s if ch.isdigit())
        return int(digits) if digits else 2
    return None
def collect_images(element, part, img_dir, safe, counter):
    """Return [(rel_path, abs_path), ...] for every image inside element's xml.
    Handles both inline drawings (a:blip/@r:embed) and legacy VML pictures
    (w:pict > v:imagedata/@r:id or w:shape/@o:relid). De-dupes by relationship id.
    """
    rids = []
    for b in element.findall('.//' + _tag('a', 'blip')):
        e = b.get(_tag('r', 'embed'))
        if e:
            rids.append(e)
    for pict in element.findall('.//' + _tag('w', 'pict')):
        for node in pict.findall('.//' + _tag('v', 'imagedata')):
            rid = node.get(_tag('r', 'id'))
            if rid:
                rids.append(rid)
        for node in pict.findall('.//' + _tag('v', 'shape')):
            rid = node.get(_tag('o', 'relid')) or node.get(_tag('r', 'id'))
            if rid:
                rids.append(rid)
    out = []
    seen = set()
    for rid in rids:
        if rid in seen:
            continue
        seen.add(rid)
        ip = part.related_parts.get(rid)
        if not ip:
            continue
        counter[0] += 1
        fn = f"img_{counter[0]:03d}.png"
        d = os.path.join(img_dir, safe)
        os.makedirs(d, exist_ok=True)
        p = os.path.join(d, fn)
        with open(p, 'wb') as f:
            f.write(ip.blob)
        out.append((f"images/{safe}/{fn}", p))
    return out
def txbx_text(element):
    """Extract text from nested text boxes (w:txbxContent)."""
    txts = []
    for tb in element.findall('.//' + _tag('w', 'txbxContent')):
        for p in tb.findall('.//' + _tag('w', 'p')):
            s = "".join(n.text or "" for n in p.findall('.//' + _tag('w', 't')))
            if s.strip():
                txts.append(s.strip())
    return "\n".join(txts)
def table_to_md(table, part, img_dir, safe, counter):
    """Emit a Markdown table.
    Uses python-docx's *normalised* cells (table.rows / row.cells), which
    already expand merged cells (gridSpan repeats the cell across the spanned
    grid columns; vMerge repeats it down the spanned rows). This keeps every
    Markdown row at the correct column count without manual span arithmetic.
    """
    out = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            txt = cell.text.strip().replace("\n", " ")
            imgs = collect_images(cell._tc, part, img_dir, safe, counter)
            if imgs:
                txt = (txt + " " + " ".join(f"[IMAGE: {r}]" for r, _ in imgs)).strip()
            cells.append(txt.replace("|", "\\|"))
        if cells:
            out.append("| " + " | ".join(cells) + " |")
    if out:
        ncols = len(out[0].split("|")) - 2  # leading/trailing pipes
        out.insert(1, "|" + "|".join(["---"] * ncols) + "|")
    return "\n".join(out)
def header_footer_md(doc, part, img_dir, safe, counter):
    """Capture default header / footer text and tables as comment blocks."""
    blocks = []
    for sec in doc.sections:
        for label, ref in (("HEADER", sec.header), ("FOOTER", sec.footer)):
            for p in ref.paragraphs:
                t = p.text.strip()
                if t:
                    blocks.append(f"<!-- {label} --> {t}")
            for tbl in ref.tables:
                md = table_to_md(tbl, part, img_dir, safe, counter)
                if md:
                    blocks.append(f"<!-- {label} TABLE -->")
                    blocks.append(md)
    return "\n".join(blocks)
def extract(docx_path, out_path=None, extract_images=False, image_dir="output/images"):
    doc = Document(docx_path)
    part = doc.part
    stem = os.path.splitext(os.path.basename(docx_path))[0]
    # Hash the source absolute path so images of two different files that share
    # a base name land in separate subdirs instead of overwriting each other.
    safe = stem + "_" + hashlib.md5(os.path.abspath(docx_path).encode("utf-8")).hexdigest()[:6]
    lines = []
    counter = [0]
    para_idx = 0
    page_idx = 1
    def emit_block(block, element=None):
        nonlocal para_idx, page_idx
        lines.append(f"\n<!-- §{para_idx} ¶约第{page_idx}页 -->")
        if element is not None:
            for br in element.findall('.//' + _tag('w', 'br')):
                if br.get(_tag('w', 'type')) == 'page':
                    page_idx += 1
        if block:
            lines.extend(block)
        para_idx += 1
    for kind, item in iter_block_items(doc):
        if kind == 'p':
            el = item._element
            txt = item.text.strip()
            tbx = txbx_text(el)
            if tbx:
                txt = (txt + "\n" + tbx).strip()
            block = []
            lvl = style_level(item)
            if lvl:
                block.append("#" * min(lvl, 6) + " " + txt)
            elif txt:
                block.append(txt)
            if extract_images:
                for r, _ in collect_images(el, part, image_dir, safe, counter):
                    block.append(f"[IMAGE: {r}]")
            emit_block(block, el)
        else:
            block = ["[TABLE]"]
            block.append(table_to_md(item, part, image_dir, safe, counter))
            block.append("[/TABLE]")
            emit_block(block, item._element)
    hf = header_footer_md(doc, part, image_dir, safe, counter)
    if hf:
        lines.append("\n" + hf)
    content = "\n".join(lines).strip() + "\n"
    if out_path:
        os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(content)
        print(f"Extracted -> {out_path} ({len(content)} chars, ~{counter[0]} images, §{para_idx} blocks)")
    else:
        print(content)
    return content
def main():
    ap = argparse.ArgumentParser(description="Extract a .docx into Markdown (V1.1).")
    ap.add_argument("input", help="input .docx path")
    ap.add_argument("output", nargs="?", help="output .md path (stdout if omitted)")
    ap.add_argument("--extract-images", action="store_true",
                    help="export images to <image-dir>/<safe>/img_NNN.png")
    ap.add_argument("--image-dir", default="output/images",
                    help="base dir for exported images (default: output/images)")
    args = ap.parse_args()
    extract(args.input, args.output, args.extract_images, args.image_dir)
if __name__ == "__main__":
    main()